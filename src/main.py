import streamlit as st
from docx import Document
import json
from agent import load_rag_system, analyze_document, check_missing_documents # Make sure to adjust the import path
from langchain_community.chat_models import ChatOpenAI
import os
import re
from docx.shared import RGBColor
from docx.text.paragraph import Paragraph
from docx.api import Document

# Set up LLM and RAG system once at the start
@st.cache_resource
def setup_agent():
    """Caches the RAG system and LLM to avoid reloading on every rerun."""
    retriever = load_rag_system()
    llm = ChatOpenAI(temperature=0, model_name="gpt-4o", openai_api_key=os.environ.get("OPENAI_API_KEY"))
    return retriever, llm

# --- UI Layout ---
st.set_page_config(layout="wide", page_title="ADGM Corporate Agent")

# Header with a professional look
st.title("ADGM Corporate Agent 💼")
st.markdown("### Intelligent AI Assistant for ADGM Document Compliance")
st.markdown("Upload your legal documents in `.docx` format for an automated review based on ADGM regulations. The agent will check for completeness, detect red flags, and provide a detailed report.")

st.markdown("---")

# File Upload Section
st.subheader("1. Document Upload")
uploaded_files = st.file_uploader(
    "Choose one or more `.docx` files to upload",
    type=['docx'],
    accept_multiple_files=True
)

if uploaded_files:
    # Use columns for a cleaner layout
    col1, col2 = st.columns([1, 2])
    
    with col1:
        st.subheader("2. Review Checklist")
        missing_docs, total_docs = check_missing_documents(uploaded_files)
        
        if missing_docs:
            st.warning(f"⚠️ **Incomplete Submission**")
            st.markdown(f"It appears you have uploaded **{len(uploaded_files)} out of {total_docs}** required documents for company incorporation.")
            st.markdown(f"**Missing Documents:** {', '.join(missing_docs)}")
            if not st.button("Proceed Anyway"):
                st.stop()
        else:
            st.success("✅ **Checklist Complete!** All mandatory documents appear to be uploaded. Starting the review process...")

    with col2:
        st.subheader("3. AI-Powered Legal Analysis")
        st.info("🔄 **Analyzing documents...** This may take a few moments. Please do not close the tab.")

        retriever, llm = setup_agent()
        full_analysis = []
        uploaded_file_objects = {f.name: f for f in uploaded_files}
        
        # Process each uploaded document
        for uploaded_file in uploaded_files:
            analysis_data = [] 
            try:
                document = Document(uploaded_file)
                doc_content = "\n".join([para.text for para in document.paragraphs])
                analysis_result_json = analyze_document(doc_content, retriever, llm)
                
                json_match = re.search(r'\[.*\]', analysis_result_json, re.DOTALL)
                if json_match:
                    json_string = json_match.group(0)
                    analysis_data = json.loads(json_string)
                else:
                    st.error(f"❌ Could not find valid JSON output for **{uploaded_file.name}**.")
                    st.text("AI Agent's raw output:")
                    st.text(analysis_result_json)
            except Exception as e:
                st.error(f"❌ Error processing **{uploaded_file.name}**: {e}")
            
            full_analysis.append({"document": uploaded_file.name, "issues": analysis_data})

        st.success("✅ **Review complete!** The analysis is ready.")

    st.markdown("---")
    
    # --- Output Section ---
    st.subheader("4. Review Report & Download")
    
    # JSON Report
    final_report = {
        "process": "Company Incorporation",
        "documents_uploaded": len(uploaded_files),
        "required_documents": total_docs,
        "missing_document": ", ".join(missing_docs) if missing_docs else "None",
        "issues_found": full_analysis
    }
    
    st.markdown("#### Structured JSON Report")
    st.json(final_report)
    
    json_str = json.dumps(final_report, indent=4)
    st.download_button(
        label="Download Full JSON Report 📄",
        data=json_str,
        file_name="adgm_analysis_report.json",
        mime="application/json"
    )

    st.markdown("---")

    # Reviewed Documents
    st.markdown("#### Reviewed Documents with Comments")
    st.info("Here you can download the reviewed documents with AI-generated comments and suggestions.")
    
    def create_reviewed_docx(uploaded_file, issues):
        document = Document(uploaded_file)
        for issue_info in issues:
            relevant_text = issue_info.get("relevant_text", "").strip()
            suggestion = issue_info.get("suggestion", "")

            if not relevant_text:
                continue

            found = False
            for paragraph in document.paragraphs:
                if found:
                    break
                if re.search(re.escape(relevant_text), paragraph.text, re.IGNORECASE):
                    comment_run = paragraph.add_run(f" [🚩 ATTENTION: {suggestion}]")
                    comment_run.bold = True
                    comment_run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                    found = True
            
            if not found:
                for table in document.tables:
                    if found:
                        break
                    for row in table.rows:
                        for cell in row.cells:
                            if re.search(re.escape(relevant_text), cell.text, re.IGNORECASE):
                                comment_paragraph = cell.add_paragraph(f" [🚩 ATTENTION: {suggestion}]")
                                comment_paragraph.runs[0].bold = True
                                comment_paragraph.runs[0].font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                                found = True
                                break
        import io
        output_stream = io.BytesIO()
        document.save(output_stream)
        return output_stream.getvalue()

    for doc_info in full_analysis:
        doc_name = doc_info['document']
        issues = doc_info['issues']
        original_file_object = uploaded_file_objects.get(doc_name)
        
        if original_file_object:
            reviewed_file_content = create_reviewed_docx(original_file_object, issues)
            st.download_button(
                label=f"Download Reviewed '{doc_name}' 📥",
                data=reviewed_file_content,
                file_name=f"reviewed_{doc_name}",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

st.markdown("---")